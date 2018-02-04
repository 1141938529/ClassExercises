import csv

import docx
import win32com.client

# 写数据到一个Word文档，text为文本内容  path-为存储地址
from openpyxl import load_workbook


def writeWordDoc(text, path):
    # 调用win32com的系统接口打开Word
    word = win32com.client.Dispatch("Word.Application")  # 操作word
    word.Visible = False  # 可以看见

    # 构造一个新的文档对象
    doc = word.Documents.Add()  # 插入文档

    # 从头部插入5行文本
    rng = doc.Range(0, 0)  # 操作位置，从00开始
    rng.InsertAfter(text)

    # 写入指定路径的文件中
    filename = path
    doc.SaveAs(filename)  # 保存
    doc.Close(True)  # Flase强行关闭,True等待缓冲区写入完毕
    # 退出Word
    word.Application.Quit()  # 退出系统接口


# 读Word中的文本内容  path--文本地址  return 段落列表
def readWordDoc(path):
    doc = docx.Document(path)
    # doc = docx.Document("demo.docx")
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return fullText


def writeExcel(dataList, path):
    # 调用系统接口打开Excel
    excel = win32com.client.Dispatch("Excel.Application")
    excel.Visible = False  # 程序窗口不可见

    # 新建工作表并定位到当前sheet页
    wb = excel.Workbooks.Add()
    currentSheet = wb.ActiveSheet
    keys = list(dataList[0].keys())
    # currentSheet.Cells(row, col).value = "Hello"
    # 第一行写入 键名
    for i in range(1, len(dataList[0]) + 1):
        currentSheet.Cells(1, i).value = keys[i - 1]
    # 开始写内容
    for row in range(2, len(dataList) + 1):
        for col in range(1, len(dataList[0]) + 1):
            currentSheet.Cells(row, col).value = dataList[row - 2][keys[col - 1]]
    # 保存内存数据到文件，关闭IO流
    wb.SaveAs(path)
    wb.Close(True)
    # 关闭Excel
    excel.Application.Quit()


def readExcel(path):
    # 载入指定路径的Excel文件到内存
    wb = load_workbook(path)

    # 拿到文件中的所有工作表
    sheetNames = wb.get_sheet_names()
    print("wb.get_sheet_names=", sheetNames)
    # 打印第一个工作表的信息
    sheet = wb.get_sheet_by_name(sheetNames[0])

    # print("sheet.title=", sheet.title)
    # print("sheet.max_row=", sheet.max_row)
    # print("sheet.max_column=", sheet.max_column)
    dataList = []
    keylist = []
    for i in range(1, 1 + sheet.max_column):
        keylist.append(sheet.cell(row=1, column=i).value)
        pass
    for row in range(2, sheet.max_row + 1):
        mDicts = {}
        for cols in range(1, 1 + sheet.max_column):
            mDicts[keylist[cols - 1]] = sheet.cell(row=row, column=cols).value

        dataList.append(mDicts)
    return dataList


def writeCsv(datalist, path):
    # 不要设置编码
    fileobj = open(path, "w", newline="")
    csvWriter = csv.writer(fileobj)
    # 打印 表头
    csvWriter.writerow(list(datalist[0].keys()))
    for itemDict in datalist:
        csvWriter.writerow(list(itemDict.values()))
    fileobj.close()
    pass


def readCsv(path):
    fileobj = open(path, "r")
    dataList = []
    alist = []
    csvReader = csv.reader(fileobj)
    for item in csvReader:
        # dataList.append(tuple(item))
        dataList.append(item)
    for i in range(1, len(dataList)):
        mdicts = {}
        # 表头为 mdicts的key值
        for j in range(len(dataList[0])):
            mdicts[dataList[0][j]] = dataList[i][j]
        alist.append(mdicts)
    return alist


if __name__ == '__main__':
    # writeWordDoc()
    msg = [
        {"name": "张三", "age": 20, "hobby": "看片"},
        {"name": "lisi", "age": 20, "hobby": "coding"},
        {"name": "wangwu", "age": 40, "hobby": "读书"},
        {"name": "zhaoliu", "age": 20, "hobby": "coding"},
    ]
    # readWordDoc()
    # writeExcel(msg, r"D:\人员信息.xlsx")
    # print(readExcel( r"D:\人员信息.xlsx"))
    print(readCsv("./res/人员信息表.csv"))
    pass
